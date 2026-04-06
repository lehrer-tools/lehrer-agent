"""
Dokument-Template für Leistungskontrollen.
Erzeugt den Grundrahmen (Kopf, Namenfeld, Notenschlüssel) -- Aufgaben werden separat eingefügt.

Verwendung:
    from template import create_test_document, add_notenschluessel, save_documents
    from omml_utils import F, T, M, add_math, add_math_line

    docs = create_test_document(
        fach="Mathematik",
        thema="Bruchrechnung",
        klassenstufe=7,
        art="Klassenarbeit",
        zeit="45 Min",
        hilfsmittel="keine",
        punkte_a=35,
        punkte_b=25,
        kurse=["A", "B"]  # oder ["A"] / ["B"]
    )

    # Aufgaben hinzufügen
    doc_a = docs["A"]
    doc_a.add_heading("Aufgabe 1: Brüche addieren (4 Punkte)", level=2)
    add_math_line(doc_a, T("a)  "), F(1, 4), T(" + "), F(2, 4), T("  = "))
    # ...

    # Notenschlüssel + Speichern
    add_notenschluessel(doc_a, 35)
    save_documents(docs, "Bruchrechnung_Kl7")
"""

import os
import tempfile

from lxml import etree
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Schulfarben
FARBE_A_KURS = RGBColor(0, 70, 140)   # Blau
FARBE_B_KURS = RGBColor(0, 120, 60)   # Grün
FARBE_IND = RGBColor(180, 90, 0)      # Orange
FARBE_SCHULE = RGBColor(128, 128, 128) # Grau

KURS_FARBEN = {"A": FARBE_A_KURS, "B": FARBE_B_KURS, "IND": FARBE_IND}
KURS_NAMEN = {
    "A": "A-Kurs (grundlegendes Niveau)",
    "B": "B-Kurs (erweitertes Niveau)",
    "IND": "Individuell",
}

# Notenschlüssel Brandenburg (VV-Leistungsbewertung Nr. 6 Abs. 3)
NOTENSCHLUESSEL = [
    (1, "sehr gut", 96),
    (2, "gut", 80),
    (3, "befriedigend", 60),
    (4, "ausreichend", 45),
    (5, "mangelhaft", 16),
    (6, "ungenügend", 0),
]

# Art -> Titel-Prefix
ART_TITEL = {
    "TÜ": "TÜ",
    "Lernerfolgskontrolle": "Lernerfolgskontrolle",
    "Klassenarbeit": "Klassenarbeit",
}


def _add_school_header(doc, schulname=None):
    """Schulname zentriert, grau, klein."""
    if not schulname:
        return None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(schulname)
    run.font.size = Pt(10)
    run.font.color.rgb = FARBE_SCHULE
    return p


def _add_header_table(doc, art, fach, thema, kurs_label=None, kurs_farbe=None):
    """Kombinierte Kopftabelle: Name/Klasse + Titel + Datum/Thema.

    Layout (2 Zeilen, 3 Spalten):
        | Name:           | Klassenarbeit    | Datum:               |
        | Klasse:         |                  | Thema: Bruchrechnung |

    Bei getrennten Kursen wird das Kurs-Label unter der Tabelle angezeigt.
    """
    prefix = ART_TITEL.get(art, art)

    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Zeile 1: Name / Titel / Datum
    r0 = table.rows[0].cells
    r0[0].text = ""
    run_name = r0[0].paragraphs[0].add_run("Name:")
    run_name.bold = True
    run_name.font.size = Pt(12)

    r0[1].text = ""
    p_title = r0[1].paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(prefix)
    run_title.bold = True
    run_title.font.size = Pt(14)

    r0[2].text = ""
    run_datum = r0[2].paragraphs[0].add_run("Datum:")
    run_datum.bold = True
    run_datum.font.size = Pt(12)

    # Zeile 2: Klasse / (leer) / Thema
    r1 = table.rows[1].cells
    r1[0].text = ""
    run_klasse = r1[0].paragraphs[0].add_run("Klasse:")
    run_klasse.bold = True
    run_klasse.font.size = Pt(12)

    r1[1].text = ""  # Mitte leer

    r1[2].text = ""
    run_thema = r1[2].paragraphs[0].add_run(f"Thema: {thema}")
    run_thema.bold = True
    run_thema.font.size = Pt(12)

    _set_table_borders(table, "000000", sz="6")

    # Kurs-Label unter der Tabelle (nur bei getrennten Dateien)
    if kurs_label and kurs_farbe:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(kurs_label)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = kurs_farbe

    return table


def _set_table_borders(table, color, sz="4", val="single"):
    """Setzt Rahmen für eine Tabelle."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else etree.SubElement(tbl, qn("w:tblPr"))
    # Vorhandene Borders entfernen
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = etree.SubElement(tblPr, qn("w:tblBorders"))
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = etree.SubElement(borders, qn(f"w:{edge}"))
        border.set(qn("w:val"), val)
        border.set(qn("w:sz"), sz)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def _remove_table_borders(table):
    """Entfernt alle Rahmen einer Tabelle."""
    _set_table_borders(table, "000000", sz="0", val="none")


def _add_separator(doc):
    """Trennlinie."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = FARBE_SCHULE
    return p


def _add_meta(doc, zeit, hilfsmittel, punkte):
    """Meta-Info (Zeit, Hilfsmittel, Punkte)."""
    text = f"Bearbeitungszeit: {zeit}"
    if hilfsmittel and hilfsmittel.lower() != "keine":
        text += f"  |  Hilfsmittel: {hilfsmittel}"
    text += f"  |  Gesamtpunktzahl: {punkte}"
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


# Rechter Tabstopp für Punktezahlen (Position in cm vom linken Rand)
_PUNKTE_TAB_CM = 16.0


def _add_right_tab(paragraph, position_cm=_PUNKTE_TAB_CM):
    """Fügt einen rechten Tabstopp zum Absatz hinzu."""
    pPr = paragraph._element.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = etree.SubElement(pPr, qn("w:tabs"))
    tab = etree.SubElement(tabs, qn("w:tab"))
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(position_cm * 567)))  # cm -> Twips


# ---------------------------------------------------------------------------
# Aufgaben-Hilfsfunktionen
# ---------------------------------------------------------------------------

def add_aufgabe(doc, nummer, titel, punkte_a, punkte_b=None):
    """Aufgaben-Überschrift mit Punkten rechtsbündig am Rand.

    Format:
        Aufgabe 1: Titel                          /3
                                              B:  /5

    Args:
        doc: docx Document
        nummer: Aufgabennummer (int)
        titel: Aufgabentitel (str)
        punkte_a: A-Kurs Punktzahl (int)
        punkte_b: B-Kurs Punktzahl (int, optional). None = keine B-Zeile.
    """
    # Zeile 1: Aufgabentext + A-Punkte rechts
    p = doc.add_paragraph()
    _add_right_tab(p)
    run = p.add_run(f"Aufgabe {nummer}:")
    run.bold = True
    run.font.size = Pt(12)
    run2 = p.add_run(f"  {titel}")
    run2.font.size = Pt(12)
    p.add_run("\t")
    run3 = p.add_run(f"_____ / {punkte_a}")
    run3.font.size = Pt(11)

    # Zeile 2: B-Punkte (optional, gleicher rechter Tabstopp)
    if punkte_b is not None:
        p2 = doc.add_paragraph()
        _add_right_tab(p2)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        p2.add_run("\t")
        r_b = p2.add_run("B  ")
        r_b.bold = True
        r_b.font.size = Pt(12)
        r_b.font.color.rgb = FARBE_B_KURS
        r_bp = p2.add_run(f"_____ / {punkte_b}")
        r_bp.font.size = Pt(11)

    return p


# Pfad zum B-Icon relativ zum Script-Verzeichnis
_B_ICON_PATH = os.path.join(os.path.dirname(__file__), "..", "assets", "b_icon.png")


def add_b_marker(doc):
    """Grünes B-Icon (Bild) als Marker für B-Kurs-Inhalte.

    Fügt das B-Icon-Bild ein. Gibt den Paragraph zurück,
    an den weiterer Text mit .add_run() angehängt werden kann.
    """
    p = doc.add_paragraph()
    icon_path = os.path.normpath(_B_ICON_PATH)
    if os.path.exists(icon_path):
        run = p.add_run()
        run.add_picture(icon_path, width=Cm(0.5))
        p.add_run(" ")
    else:
        # Fallback: grüner Text
        run = p.add_run("B ")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = FARBE_B_KURS
    return p


def add_rechenkaestchen(doc, rows=5, cols=20, cell_size_cm=0.55):
    """Rechenkästchen (kariertes Gitter) als Antwortfläche.

    Erzeugt eine Tabelle mit quadratischen Zellen, die wie kariertes Papier
    aussehen -- so wie in den echten Klassenarbeiten der Schule.

    Args:
        doc: docx Document
        rows: Anzahl Zeilen (Standard: 5)
        cols: Anzahl Spalten (Standard: 20)
        cell_size_cm: Kantenlänge einer Zelle in cm (Standard: 0.55)
    """
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    size = Cm(cell_size_cm)
    for row in table.rows:
        row.height = size
        for cell in row.cells:
            cell.width = size
            # Innenabstand auf 0
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcMar = tcPr.find(qn("w:tcMar"))
            if tcMar is None:
                tcMar = etree.SubElement(tcPr, qn("w:tcMar"))
            for side in ("top", "bottom", "start", "end"):
                el = etree.SubElement(tcMar, qn(f"w:{side}"))
                el.set(qn("w:w"), "0")
                el.set(qn("w:type"), "dxa")
            # Zellentext leer, Absatzabstände minimieren
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                # Font-Größe auf Run-Ebene, NICHT auf Style-Ebene
                run = p.add_run("")
                run.font.size = Pt(2)

    _set_table_borders(table, "999999", sz="4")
    doc.add_paragraph()  # Abstand
    return table


def add_answer_lines(doc, count=3):
    """Leere Antwortzeilen für Textantworten.

    Args:
        doc: docx Document
        count: Anzahl Zeilen (Standard: 3)
    """
    for _ in range(count):
        p = doc.add_paragraph()
        run = p.add_run("_" * 80)
        run.font.size = Pt(10)
        run.font.color.rgb = FARBE_SCHULE
        pf = p.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(0)
    doc.add_paragraph()
    return


def add_value_table(doc, headers, data):
    """Wertetabelle mit optionalen Leerfeldern.

    Args:
        doc: docx Document
        headers: Liste von Spaltenüberschriften, z.B. ["x", "y"]
        data: Liste von Zeilen, z.B. [[0, ""], [1, ""], ["", 3]]
              Leere Strings werden als auszufüllende Felder dargestellt.
    """
    cols = len(headers)
    table = doc.add_table(rows=1 + len(data), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        run.bold = True
        run.font.size = Pt(11)

    # Daten
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val) if val != "" else "")
            run.font.size = Pt(11)

    _set_table_borders(table, "000000", sz="6")
    doc.add_paragraph()
    return table


def add_coordinate_system(doc, x_range=(-5, 5), y_range=(-5, 5),
                          width_cm=12, height_cm=12, grid=True):
    """Leeres Koordinatensystem als Bild einfügen.

    Benötigt matplotlib. Erzeugt ein temporäres PNG und fügt es ein.

    Args:
        doc: docx Document
        x_range: Tuple (min, max) für x-Achse
        y_range: Tuple (min, max) für y-Achse
        width_cm: Breite in cm
        height_cm: Höhe in cm
        grid: Gitternetz anzeigen
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(1, 1, figsize=(width_cm / 2.54, height_cm / 2.54))

    ax.set_xlim(x_range)
    ax.set_ylim(y_range)
    ax.set_aspect("equal")

    # Achsen durch den Ursprung
    ax.spines["left"].set_position("zero")
    ax.spines["bottom"].set_position("zero")
    ax.spines["right"].set_visible(False)
    ax.spines["top"].set_visible(False)

    # Pfeile
    ax.plot(1, 0, ">k", transform=ax.get_yaxis_transform(), clip_on=False)
    ax.plot(0, 1, "^k", transform=ax.get_xaxis_transform(), clip_on=False)

    # Beschriftung
    ax.set_xlabel("x", fontsize=12, loc="right")
    ax.set_ylabel("y", fontsize=12, loc="top", rotation=0)

    # Ganzzahlige Ticks
    import numpy as np
    ax.set_xticks(np.arange(x_range[0], x_range[1] + 1, 1))
    ax.set_yticks(np.arange(y_range[0], y_range[1] + 1, 1))
    ax.tick_params(labelsize=8)

    # 0 nicht doppelt anzeigen
    xlabels = [str(int(t)) if t != 0 else "" for t in ax.get_xticks()]
    ylabels = [str(int(t)) if t != 0 else "" for t in ax.get_yticks()]
    ax.set_xticklabels(xlabels)
    ax.set_yticklabels(ylabels)

    if grid:
        ax.grid(True, linewidth=0.3, color="gray", alpha=0.5)

    plt.tight_layout()

    # Temporäre Datei
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmp.name, dpi=150, bbox_inches="tight")
    plt.close(fig)

    doc.add_picture(tmp.name, width=Cm(width_cm))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    os.unlink(tmp.name)
    return last_paragraph


def add_number_line(doc, start=0, end=2, divisions=10, width_cm=14,
                    labels=None, points=None):
    """Zahlenstrahl als Bild einfügen.

    Args:
        doc: docx Document
        start: Startwert (z.B. 0)
        end: Endwert (z.B. 2)
        divisions: Unterteilungen pro ganzer Einheit (z.B. 10 für Zehntel)
        width_cm: Breite in cm
        labels: dict {position: "Label"} für Beschriftungen unter dem Strahl
        points: dict {position: "Buchstabe"} für markierte Punkte über dem Strahl
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np

    fig, ax = plt.subplots(1, 1, figsize=(width_cm / 2.54, 2.0))

    # Strahl
    ax.arrow(start, 0, end - start + 0.3, 0,
             head_width=0.06, head_length=0.05,
             fc="black", ec="black", linewidth=1.5)

    # Teilstriche
    total_ticks = int((end - start) * divisions) + 1
    for i in range(total_ticks):
        x = start + i / divisions
        if i % divisions == 0:
            # Ganzzahlige Markierung (lang)
            ax.plot([x, x], [-0.08, 0.08], "k-", linewidth=1.5)
            ax.text(x, -0.18, str(int(x)), ha="center", va="top",
                    fontsize=11, fontweight="bold")
        else:
            # Zwischenmarkierung (kurz)
            ax.plot([x, x], [-0.04, 0.04], "k-", linewidth=0.8)

    # Beschriftungen unter dem Strahl
    if labels:
        for pos, label in labels.items():
            ax.text(pos, -0.3, label, ha="center", va="top", fontsize=9)

    # Markierte Punkte über dem Strahl
    if points:
        for pos, letter in points.items():
            ax.plot([pos, pos], [0, 0.12], "k-", linewidth=1.2)
            ax.text(pos, 0.16, letter, ha="center", va="bottom",
                    fontsize=11, fontweight="bold")

    ax.set_xlim(start - 0.1, end + 0.5)
    ax.set_ylim(-0.4, 0.3)
    ax.set_aspect("auto")
    ax.axis("off")
    plt.tight_layout(pad=0.1)

    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmp.name, dpi=150, bbox_inches="tight")
    plt.close(fig)

    doc.add_picture(tmp.name, width=Cm(width_cm))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    os.unlink(tmp.name)
    return last_paragraph


def add_bewertungsfuss(doc, gesamtpunkte, punkte_b=None, art="Klassenarbeit"):
    """Bewertungsfuß am Ende des Dokuments.

    Layout:
        ──────────────────────────
        Punkte:       / 39     Schulleitung:        (nur bei Klassenarbeit)
                 [B]  / 52
        Zensur:                Unterschrift der ...: (nur bei Klassenarbeit)

    Args:
        doc: docx Document
        gesamtpunkte: Gesamtpunktzahl dieses Kurses (int)
        punkte_b: B-Kurs-Punktzahl (grünes B-Icon). None = kein B-Feld.
        art: Art der Arbeit. Schulleitung + Unterschrift nur bei "Klassenarbeit".
    """
    ist_klassenarbeit = art == "Klassenarbeit"
    _add_separator(doc)

    # Zeilenanzahl berechnen
    num_rows = 1  # Punkte-Zeile
    if punkte_b is not None:
        num_rows += 1  # B-Zeile
    num_rows += 1  # Leerzeile
    num_rows += 1  # Zensur-Zeile

    table = doc.add_table(rows=num_rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    row_idx = 0
    icon_path = os.path.normpath(_B_ICON_PATH)

    # Zeile 1: Punkte + ggf. Schulleitung
    c_punkte = table.rows[row_idx].cells[0]
    c_right = table.rows[row_idx].cells[1]

    c_punkte.text = ""
    p1 = c_punkte.paragraphs[0]
    r1a = p1.add_run("Punkte:")
    r1a.bold = True
    r1a.font.size = Pt(12)
    r1b = p1.add_run(f"              / {gesamtpunkte}")
    r1b.font.size = Pt(12)

    c_right.text = ""
    if ist_klassenarbeit:
        p1r = c_right.paragraphs[0]
        r1c = p1r.add_run("Schulleitung:")
        r1c.bold = True
        r1c.font.size = Pt(12)

    row_idx += 1

    # Zeile 2: B-Punkte mit Icon (optional)
    if punkte_b is not None:
        c_b = table.rows[row_idx].cells[0]
        c_b.text = ""
        p2 = c_b.paragraphs[0]
        p2.add_run("          ").font.size = Pt(12)
        if os.path.exists(icon_path):
            run_icon = p2.add_run()
            run_icon.add_picture(icon_path, width=Cm(0.5))
        else:
            r2b = p2.add_run("B")
            r2b.bold = True
            r2b.font.size = Pt(14)
            r2b.font.color.rgb = FARBE_B_KURS
        r2c = p2.add_run(f"          / {punkte_b}")
        r2c.font.size = Pt(12)
        table.rows[row_idx].cells[1].text = ""
        row_idx += 1

    # Leerzeile
    table.rows[row_idx].cells[0].text = ""
    table.rows[row_idx].cells[1].text = ""
    row_idx += 1

    # Letzte Zeile: Zensur + ggf. Unterschrift
    c_zensur = table.rows[row_idx].cells[0]
    c_unt = table.rows[row_idx].cells[1]

    c_zensur.text = ""
    p3 = c_zensur.paragraphs[0]
    r3a = p3.add_run("Zensur:")
    r3a.bold = True
    r3a.font.size = Pt(12)

    c_unt.text = ""
    if ist_klassenarbeit:
        p4 = c_unt.paragraphs[0]
        r4a = p4.add_run("Unterschrift der Erziehungsberechtigten:")
        r4a.bold = True
        r4a.font.size = Pt(12)

    _remove_table_borders(table)
    return table


def add_notenschluessel(doc, gesamtpunkte):
    """Notenschlüssel-Tabelle am Ende des Dokuments."""
    _add_separator(doc)
    p = doc.add_paragraph()
    run = p.add_run("Notenschlüssel")
    run.bold = True
    run.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    header[0].text = "Note"
    header[1].text = "Bezeichnung"
    header[2].text = "ab %"
    header[3].text = f"ab Punkte (von {gesamtpunkte})"
    for cell in header:
        for p in cell.paragraphs:
            run = p.runs[0] if p.runs else p.add_run()
            run.bold = True
            run.font.size = Pt(9)

    for note, bezeichnung, prozent in NOTENSCHLUESSEL:
        row = table.add_row().cells
        row[0].text = str(note)
        row[1].text = bezeichnung
        row[2].text = f"{prozent}%"
        punkte = int(gesamtpunkte * prozent / 100 + 0.5) if prozent > 0 else 0
        row[3].text = str(punkte)
        for cell in row:
            for p in cell.paragraphs:
                if p.runs:
                    p.runs[0].font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Erreichte Punktzahl: _____ / ")
    run.font.size = Pt(11)
    run2 = p.add_run(str(gesamtpunkte))
    run2.bold = True
    run2.font.size = Pt(11)
    p.add_run("          Note: _____").font.size = Pt(11)
    return table


def create_test_document(fach, thema, klassenstufe, art, zeit, hilfsmittel,
                         punkte_a=None, punkte_b=None, punkte_ind=None,
                         kurse=None, ind_schueler=None, modus="getrennt"):
    """Erzeugt Dokument(e) für Leistungskontrollen.

    Args:
        fach: z.B. "Mathematik"
        thema: z.B. "Bruchrechnung"
        klassenstufe: z.B. 7
        art: "Kurze Kontrolle", "Test", "Klassenarbeit", "Große Klassenarbeit"
        zeit: z.B. "45 Min"
        hilfsmittel: z.B. "keine" oder "Taschenrechner"
        punkte_a: Gesamtpunktzahl A-Kurs
        punkte_b: Gesamtpunktzahl B-Kurs
        punkte_ind: Gesamtpunktzahl individuelle Variante
        kurse: Liste, z.B. ["A", "B"] oder ["A", "B", "IND"]
        ind_schueler: Name(n) für individuelle Variante, z.B. "Ory und Layla"
        modus: "getrennt" = separate Dateien pro Kurs (Standard)
               "integriert" = EIN Dokument, B-Aufgaben mit grünem B markiert

    Returns:
        modus="getrennt": dict {"A": Document, "B": Document, ...}
        modus="integriert": dict {"AB": Document} (ggf. + "IND": Document)
    """
    if kurse is None:
        kurse = ["A", "B"]

    punkte = {"A": punkte_a, "B": punkte_b, "IND": punkte_ind}
    docs = {}

    if modus == "integriert":
        # Ein Dokument für A+B zusammen
        doc = _create_doc_skeleton(art, fach, thema, zeit, hilfsmittel, None)
        docs["AB"] = doc

        # IND trotzdem als eigene Datei
        if "IND" in kurse and ind_schueler:
            doc_ind = _create_doc_skeleton(
                art, fach, thema, zeit, hilfsmittel, punkte.get("IND"),
                kurs_label=ind_schueler, kurs_farbe=FARBE_IND,
            )
            docs["IND"] = doc_ind
    else:
        # Getrennte Dateien pro Kurs
        for kurs in kurse:
            label = None
            farbe = None
            if kurs == "IND" and ind_schueler:
                label = ind_schueler
                farbe = FARBE_IND
            elif kurs in KURS_NAMEN:
                label = KURS_NAMEN[kurs]
                farbe = KURS_FARBEN[kurs]
            doc = _create_doc_skeleton(
                art, fach, thema, zeit, hilfsmittel, punkte.get(kurs),
                kurs_label=label, kurs_farbe=farbe,
            )
            docs[kurs] = doc

    return docs


def _create_doc_skeleton(art, fach, thema, zeit, hilfsmittel, punkte_total,
                         kurs_label=None, kurs_farbe=None):
    """Erzeugt ein einzelnes Dokument mit Grundgerüst."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Calibri"

    # Kombinierte Kopftabelle (Name/Klasse + Titel + Datum/Thema)
    _add_header_table(doc, art, fach, thema,
                      kurs_label=kurs_label, kurs_farbe=kurs_farbe)

    doc.add_paragraph()
    return doc


def save_documents(docs, basename):
    """Speichert die Dokumente als separate Dateien pro Kurs.

    Args:
        docs: dict {"A": Document, ...} von create_test_document()
        basename: z.B. "Bruchrechnung_Kl7"

    Returns:
        Liste der erzeugten Dateipfade
    """
    paths = []
    kurs_labels = {
        "A": "A-Kurs", "B": "B-Kurs",
        "AB": "integriert", "IND": "Individuell",
    }
    for kurs, doc in docs.items():
        label = kurs_labels.get(kurs, kurs)
        filename = f"{basename}_{label}.docx"
        doc.save(filename)
        paths.append(filename)
        print(f"Gespeichert: {filename}")
    return paths
