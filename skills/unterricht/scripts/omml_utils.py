"""
OMML-Hilfsfunktionen für echte Word-Formeln (Brüche, gemischte Zahlen, Operatoren).

Verwendung:
    from omml_utils import add_math, F, T, M, add_math_line

    # Bruch inline
    p = doc.add_paragraph()
    p.add_run("Berechne ")
    add_math(p, F(2, 3), T(" + "), F(1, 4))

    # Formelzeile als eigener Absatz
    add_math_line(doc, T("a)  "), F(2, 3), T(" + "), F(1, 4), T("  = "))
"""

from lxml import etree
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Unicode-Operatoren
MINUS = "\u2212"
MAL = "\u00b7"
GETEILT = "\u00f7"
KREUZ = "\u00d7"


def _math_run(parent, text):
    """Erzeugt einen m:r-Lauf in Cambria Math."""
    r = etree.SubElement(parent, qn("m:r"))
    wRPr = etree.SubElement(r, qn("w:rPr"))
    rF = etree.SubElement(wRPr, qn("w:rFonts"))
    rF.set(qn("w:ascii"), "Cambria Math")
    rF.set(qn("w:hAnsi"), "Cambria Math")
    t = etree.SubElement(r, qn("m:t"))
    t.text = str(text)
    return r


def _math_text_run(parent, text):
    """Erzeugt einen normalen (nicht-kursiven) Text-Lauf innerhalb einer Formel."""
    r = etree.SubElement(parent, qn("m:r"))
    mRPr = etree.SubElement(r, qn("m:rPr"))
    etree.SubElement(mRPr, qn("m:nor"))
    wRPr = etree.SubElement(r, qn("w:rPr"))
    rF = etree.SubElement(wRPr, qn("w:rFonts"))
    rF.set(qn("w:ascii"), "Cambria Math")
    rF.set(qn("w:hAnsi"), "Cambria Math")
    t = etree.SubElement(r, qn("m:t"))
    t.set(qn("xml:space"), "preserve")
    t.text = text
    return r


def _build_fraction(parent, zaehler, nenner):
    """Erzeugt einen OMML-Bruch (m:f) mit Bruchstrich."""
    f = etree.SubElement(parent, qn("m:f"))
    fPr = etree.SubElement(f, qn("m:fPr"))
    ftype = etree.SubElement(fPr, qn("m:type"))
    ftype.set(qn("m:val"), "bar")
    num = etree.SubElement(f, qn("m:num"))
    _math_run(num, zaehler)
    den = etree.SubElement(f, qn("m:den"))
    _math_run(den, nenner)
    return f


# Leerfeld-Zeichen für auszufüllende Brüche
LEER = "\u25a1"  # □ weißes Quadrat


# --- Öffentliche Element-Konstruktoren ---

def F(zaehler, nenner):
    """Bruch: F(2, 3) -> 2/3, F(LEER, LEER) -> □/□"""
    return ("frac", zaehler, nenner)


def T(text):
    """Text/Operator in Formel: T(" + ")"""
    return ("text", text)


def M(ganze, zaehler, nenner):
    """Gemischte Zahl: M(2, 1, 3) -> 2 1/3, M(LEER, LEER, LEER) -> □ □/□"""
    return ("mixed", ganze, zaehler, nenner)


# --- Rendering ---

def add_math(paragraph, *elements):
    """Fügt OMML-Formelelemente in einen bestehenden Absatz ein.

    Args:
        paragraph: docx Paragraph-Objekt
        *elements: Ergebnis von F(), T(), M() Aufrufen
    """
    omath = etree.SubElement(paragraph._element, qn("m:oMath"))

    for elem in elements:
        kind = elem[0]
        if kind == "frac":
            _, z, n = elem
            _build_fraction(omath, z, n)
        elif kind == "text":
            _, txt = elem
            _math_text_run(omath, txt)
        elif kind == "mixed":
            _, ganze, z, n = elem
            _math_run(omath, ganze)
            _build_fraction(omath, z, n)


def add_math_line(doc, *elements, alignment="left"):
    """Erzeugt einen neuen Absatz mit Formelelementen (linksbündig).

    Nutzt m:oMathPara mit m:jc für echte OMML-Ausrichtung.

    Args:
        doc: docx Document-Objekt
        *elements: Ergebnis von F(), T(), M() Aufrufen
        alignment: "left" (Standard), "center", "right"

    Returns:
        Der erzeugte Paragraph
    """
    p = doc.add_paragraph()

    # OMML-Formelabsatz mit Ausrichtung
    oMathPara = etree.SubElement(p._element, qn("m:oMathPara"))
    oMathParaPr = etree.SubElement(oMathPara, qn("m:oMathParaPr"))
    jc = etree.SubElement(oMathParaPr, qn("m:jc"))
    jc.set(qn("m:val"), alignment)

    omath = etree.SubElement(oMathPara, qn("m:oMath"))

    for elem in elements:
        kind = elem[0]
        if kind == "frac":
            _, z, n = elem
            _build_fraction(omath, z, n)
        elif kind == "text":
            _, txt = elem
            _math_text_run(omath, txt)
        elif kind == "mixed":
            _, ganze, z, n = elem
            _math_run(omath, ganze)
            _build_fraction(omath, z, n)

    return p
