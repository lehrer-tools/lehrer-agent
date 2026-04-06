"""
Dokument-Template für Unterrichtsvorbereitungen (Hefteintrag/Tafelbild).
Erzeugt .docx-Dateien mit Merke-Kästen, Methoden-Schritten, Beispielen und Aufgabenverweisen.

Verwendung:
    from template import (
        create_lesson_document, add_merke, add_methode,
        add_beispiel, add_probe, add_aufgaben, add_zusatzmaterial,
        add_coordinate_system, add_value_table, save_document
    )

    doc = create_lesson_document(
        thema="Einsetzungsverfahren",
        kapitel="3 — Lineare Gleichungssysteme",
        seiten="S. 91-92"
    )
    add_merke(doc, "Beim Einsetzungsverfahren löst man ...")
    add_methode(doc, "Einsetzungsverfahren", ["Schritt 1", "Schritt 2"])
    add_beispiel(doc, "Beispiel")
    add_probe(doc)
    add_aufgaben(doc, ["B.S. 92/1a-c"], ["AH S. 35/1,2"])
    save_document(doc, "Einsetzungsverfahren")
"""

import os
import tempfile

from lxml import etree
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Farben
FARBE_TITEL = RGBColor(0, 70, 140)      # Blau — Titel
FARBE_MERKE = RGBColor(0, 70, 140)      # Blau — Merke-Rahmen
FARBE_MERKE_BG = "E8F0FE"               # Hellblau — Merke-Hintergrund
FARBE_METHODE = RGBColor(0, 100, 60)    # Grün — Methode
FARBE_BEISPIEL = RGBColor(80, 80, 80)   # Dunkelgrau — Beispiel
FARBE_GRAU = RGBColor(128, 128, 128)    # Grau — Untertitel


def _set_cell_shading(cell, color_hex):
    """Setzt Hintergrundfarbe einer Tabellenzelle."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shading = etree.SubElement(tcPr, qn("w:shd"))
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)


def _set_table_borders(table, color, sz="4", val="single"):
    """Setzt Rahmen für eine Tabelle."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else etree.SubElement(tbl, qn("w:tblPr"))
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = etree.SubElement(tblPr, qn("w:tblBorders"))
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = etree.SubElement(borders, qn(f"w:{edge}"))
        border.set(qn("w:val"), val)
        border.set(qn("w:sz"), sz)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def _remove_table_borders(table):
    """Entfernt alle Rahmen einer Tabelle."""
    _set_table_borders(table, "000000", sz="0", val="none")


def _add_separator(doc):
    """Dünne Trennlinie."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("─" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = FARBE_GRAU


# ---------------------------------------------------------------------------
# Dokument erstellen
# ---------------------------------------------------------------------------

def create_lesson_document(thema, kapitel=None, seiten=None):
    """Erzeugt ein neues Unterrichtsdokument mit Kopfbereich.

    Args:
        thema: Thema der Stunde, z.B. "Einsetzungsverfahren"
        kapitel: Kapitelangabe, z.B. "3 — Lineare Gleichungssysteme"
        seiten: Seitenangabe, z.B. "S. 91-92"

    Returns:
        docx Document
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Calibri"

    # Titel
    p = doc.add_paragraph()
    run = p.add_run(thema)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = FARBE_TITEL

    # Untertitel: Kapitel + Seiten
    if kapitel or seiten:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        parts = []
        if kapitel:
            parts.append(f"Kapitel {kapitel}")
        if seiten:
            parts.append(seiten)
        run2 = p2.add_run(" | ".join(parts))
        run2.font.size = Pt(10)
        run2.font.color.rgb = FARBE_GRAU

    return doc


# ---------------------------------------------------------------------------
# Merke-Kasten
# ---------------------------------------------------------------------------

def add_merke(doc, text):
    """Hervorgehobener Merke-Kasten mit blauem Rahmen und hellblauem Hintergrund.

    Args:
        doc: docx Document
        text: Inhalt des Merke-Kastens (kann mehrzeilig sein, \\n wird zu Absätzen)
    """
    # Überschrift
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Merke")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = FARBE_MERKE

    # Kasten als 1-Zellen-Tabelle
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Breite auf volle Seitenbreite
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else etree.SubElement(tbl, qn("w:tblPr"))
    tblW = etree.SubElement(tblPr, qn("w:tblW"))
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")  # 100%

    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, FARBE_MERKE_BG)

    # Text einfügen
    cell.text = ""
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        run = p.add_run(line.strip())
        run.font.size = Pt(12)

    _set_table_borders(table, "0046C8", sz="8")  # Blauer Rahmen

    doc.add_paragraph()  # Abstand
    return table


# ---------------------------------------------------------------------------
# Methode / Schritte
# ---------------------------------------------------------------------------

def add_methode(doc, title, steps):
    """Nummerierte Methodenschritte.

    Args:
        doc: docx Document
        title: Name der Methode, z.B. "Einsetzungsverfahren"
        steps: Liste von Schritten (Strings)
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = FARBE_METHODE

    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run_num = p.add_run(f"{i}. ")
        run_num.bold = True
        run_num.font.size = Pt(12)
        run_text = p.add_run(step)
        run_text.font.size = Pt(12)

    doc.add_paragraph()  # Abstand


# ---------------------------------------------------------------------------
# Beispiel / Probe
# ---------------------------------------------------------------------------

def add_beispiel(doc, title="Beispiel"):
    """Beispiel-Überschrift. Danach Formeln/Text mit add_math_line oder doc.add_paragraph einfügen.

    Args:
        doc: docx Document
        title: Überschrift, Standard "Beispiel"

    Returns:
        Der Überschrift-Paragraph
    """
    _add_separator(doc)
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = FARBE_BEISPIEL
    return p


def add_probe(doc):
    """Probe-Überschrift. Danach Formeln/Text einfügen.

    Returns:
        Der Überschrift-Paragraph
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Probe:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = FARBE_BEISPIEL
    return p


# ---------------------------------------------------------------------------
# Aufgaben
# ---------------------------------------------------------------------------

def _render_aufgaben_liste(doc, items):
    """Rendert eine Liste von Aufgaben (Strings oder Dicts mit ref+text)."""
    for item in items:
        if isinstance(item, dict):
            # Dict: {"ref": "B.S. 92/1a-c", "text": "Berechne den Flächeninhalt..."}
            ref = item.get("ref", "")
            text = item.get("text", "")
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(ref)
            run.bold = True
            run.font.size = Pt(11)
            if text:
                for line in text.split("\n"):
                    tp = doc.add_paragraph()
                    tp.paragraph_format.space_before = Pt(0)
                    tp.paragraph_format.space_after = Pt(1)
                    tp.paragraph_format.left_indent = Cm(0.5)
                    run = tp.add_run(line)
                    run.font.size = Pt(11)
        else:
            # Einfacher String: nur Referenz
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            run = p.add_run(str(item))
            run.font.size = Pt(12)


def add_aufgaben(doc, unterricht=None, hausaufgabe=None):
    """Aufgabenverweise für Unterricht und Hausaufgabe.

    Jeder Eintrag kann ein einfacher String sein (nur Referenz)
    oder ein Dict mit Aufgabentext:

        # Nur Referenz:
        add_aufgaben(doc, ["B.S. 92/1a-c"])

        # Mit Text:
        add_aufgaben(doc, [{"ref": "B.S. 92/1a-c", "text": "Berechne den Flächeninhalt..."}])

        # Gemischt:
        add_aufgaben(doc, ["B.S. 92/1", {"ref": "B.S. 92/2", "text": "Zeichne die Gerade..."}])

    Args:
        doc: docx Document
        unterricht: String, Dict oder Liste davon
        hausaufgabe: String, Dict oder Liste davon
    """
    # Einzelwerte in Listen umwandeln
    if isinstance(unterricht, (str, dict)):
        unterricht = [unterricht]
    if isinstance(hausaufgabe, (str, dict)):
        hausaufgabe = [hausaufgabe]

    _add_separator(doc)
    p = doc.add_paragraph()
    run = p.add_run("Aufgaben")
    run.bold = True
    run.font.size = Pt(13)

    if unterricht:
        p = doc.add_paragraph()
        run = p.add_run("Im Unterricht:")
        run.bold = True
        run.font.size = Pt(12)
        _render_aufgaben_liste(doc, unterricht)

    if hausaufgabe:
        p = doc.add_paragraph()
        run = p.add_run("Hausaufgabe:")
        run.bold = True
        run.font.size = Pt(12)
        _render_aufgaben_liste(doc, hausaufgabe)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Zusatzmaterial
# ---------------------------------------------------------------------------

def add_zusatzmaterial(doc, materials):
    """Liste relevanter Zusatzmaterialien mit Inhalt.

    Args:
        doc: docx Document
        materials: Liste von Tupeln (titel, dateiname) oder (titel, dateiname, books_dir), z.B.
                   [("Überblick aller Lösungsverfahren", "3507_81278_051.doc")]
                   Wenn books_dir angegeben, wird der Inhalt aus
                   books_dir/Zusatzmaterial/<dateiname>.md eingelesen und eingebettet.
    """
    if not materials:
        return

    p = doc.add_paragraph()
    run = p.add_run("Zusatzmaterial")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = FARBE_GRAU

    for item in materials:
        titel = item[0]
        datei = item[1]
        books_dir = item[2] if len(item) > 2 else None

        # Überschrift
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"\u2022  {titel}")
        run.bold = True
        run.font.size = Pt(11)

        # Inhalt einlesen falls Pfad angegeben
        if books_dir:
            # Dateiname ohne Endung + .md
            md_name = os.path.splitext(datei)[0] + ".md"
            md_path = os.path.join(books_dir, "Zusatzmaterial", md_name)
            if os.path.exists(md_path):
                with open(md_path, "r", encoding="utf-8") as f:
                    content = f.read().strip()
                if content:
                    for line in content.split("\n"):
                        cp = doc.add_paragraph()
                        cp.paragraph_format.space_before = Pt(0)
                        cp.paragraph_format.space_after = Pt(1)
                        run = cp.add_run(line)
                        run.font.size = Pt(10)
            else:
                p2 = doc.add_paragraph()
                run2 = p2.add_run(f"  ({datei})")
                run2.font.size = Pt(9)
                run2.font.color.rgb = FARBE_GRAU
        else:
            p2 = doc.add_paragraph()
            run2 = p2.add_run(f"  ({datei})")
            run2.font.size = Pt(9)
            run2.font.color.rgb = FARBE_GRAU


# ---------------------------------------------------------------------------
# Koordinatensystem
# ---------------------------------------------------------------------------

def add_coordinate_system(doc, x_range=(-5, 5), y_range=(-5, 5),
                          width_cm=12, height_cm=12, grid=True):
    """Leeres Koordinatensystem als Bild einfügen.

    Args:
        doc: docx Document
        x_range: Tuple (min, max) für x-Achse
        y_range: Tuple (min, max) für y-Achse
        width_cm: Breite in cm
        height_cm: Höhe in cm
        grid: Gitternetz anzeigen
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np

    fig, ax = plt.subplots(1, 1, figsize=(width_cm / 2.54, height_cm / 2.54))

    ax.set_xlim(x_range)
    ax.set_ylim(y_range)
    ax.set_aspect("equal")

    # Achsen durch den Ursprung
    ax.spines["left"].set_position("zero")
    ax.spines["bottom"].set_position("zero")
    ax.spines["right"].set_visible(False)
    ax.spines["top"].set_visible(False)

    # Pfeile
    ax.plot(1, 0, ">k", transform=ax.get_yaxis_transform(), clip_on=False)
    ax.plot(0, 1, "^k", transform=ax.get_xaxis_transform(), clip_on=False)

    # Beschriftung
    ax.set_xlabel("x", fontsize=12, loc="right")
    ax.set_ylabel("y", fontsize=12, loc="top", rotation=0)

    # Ganzzahlige Ticks
    ax.set_xticks(np.arange(x_range[0], x_range[1] + 1, 1))
    ax.set_yticks(np.arange(y_range[0], y_range[1] + 1, 1))
    ax.tick_params(labelsize=8)

    # 0 nicht doppelt anzeigen
    xlabels = [str(int(t)) if t != 0 else "" for t in ax.get_xticks()]
    ylabels = [str(int(t)) if t != 0 else "" for t in ax.get_yticks()]
    ax.set_xticklabels(xlabels)
    ax.set_yticklabels(ylabels)

    if grid:
        ax.grid(True, linewidth=0.3, color="gray", alpha=0.5)

    plt.tight_layout()

    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmp.name, dpi=150, bbox_inches="tight")
    plt.close(fig)

    doc.add_picture(tmp.name, width=Cm(width_cm))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    os.unlink(tmp.name)
    return last_paragraph


# ---------------------------------------------------------------------------
# Wertetabelle
# ---------------------------------------------------------------------------

def add_value_table(doc, headers, data):
    """Wertetabelle.

    Args:
        doc: docx Document
        headers: Liste von Spaltenüberschriften, z.B. ["x", "y"]
        data: Liste von Zeilen, z.B. [[0, 3], [1, 5], [2, 7]]
    """
    cols = len(headers)
    table = doc.add_table(rows=1 + len(data), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        run.bold = True
        run.font.size = Pt(11)

    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val) if val != "" else "")
            run.font.size = Pt(11)

    _set_table_borders(table, "000000", sz="6")
    doc.add_paragraph()
    return table


# ---------------------------------------------------------------------------
# Lösungen / Erwartungshorizont
# ---------------------------------------------------------------------------

def add_loesungen(doc, loesungen):
    """Lösungsblatt für den Lehrer (Seitenumbruch davor).

    Args:
        doc: docx Document
        loesungen: Liste von Dicts mit ref + lösung, z.B.
                   [{"ref": "B.S. 92/1a", "loesung": "x = 3, y = 5"}]
                   oder Liste von Tupeln: [("B.S. 92/1a", "x = 3, y = 5")]
    """
    if not loesungen:
        return

    # Seitenumbruch
    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run("Lösungen / Erwartungshorizont")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = FARBE_GRAU

    p2 = doc.add_paragraph()
    run2 = p2.add_run("(nur für die Lehrkraft)")
    run2.font.size = Pt(9)
    run2.font.italic = True
    run2.font.color.rgb = FARBE_GRAU

    for item in loesungen:
        if isinstance(item, dict):
            ref = item.get("ref", "")
            loesung = item.get("loesung", "")
        else:
            ref, loesung = item[0], item[1]

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(ref)
        run.bold = True
        run.font.size = Pt(11)

        if loesung:
            for line in str(loesung).split("\n"):
                lp = doc.add_paragraph()
                lp.paragraph_format.space_before = Pt(0)
                lp.paragraph_format.space_after = Pt(1)
                lp.paragraph_format.left_indent = Cm(0.5)
                run = lp.add_run(line)
                run.font.size = Pt(11)


# ---------------------------------------------------------------------------
# Stundenkontext
# ---------------------------------------------------------------------------

def add_stundenkontext(doc, vorher=None, nachher=None):
    """Kontext-Leiste am Dokumentende: was kam vorher, was kommt danach.

    Args:
        doc: docx Document
        vorher: String, z.B. "Lineare Funktionen (S. 78-88)"
        nachher: String, z.B. "Scheitelpunktform (S. 105-110)"
    """
    if not vorher and not nachher:
        return

    _add_separator(doc)
    table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(table)

    if vorher:
        cell = table.rows[0].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run("\u2190 Vorher: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = FARBE_GRAU
        run2 = p.add_run(vorher)
        run2.font.size = Pt(9)
        run2.font.color.rgb = FARBE_GRAU

    if nachher:
        cell = table.rows[0].cells[1]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("Nachher: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = FARBE_GRAU
        run2 = p.add_run(nachher + " \u2192")
        run2.font.size = Pt(9)
        run2.font.color.rgb = FARBE_GRAU


# ---------------------------------------------------------------------------
# Speichern
# ---------------------------------------------------------------------------

def save_document(doc, basename):
    """Speichert das Dokument als .docx.

    Args:
        doc: docx Document
        basename: Dateiname ohne Endung, z.B. "Einsetzungsverfahren"

    Returns:
        Dateipfad
    """
    filename = f"{basename}.docx"
    doc.save(filename)
    print(f"Gespeichert: {filename}")
    return filename
